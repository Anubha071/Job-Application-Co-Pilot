import os
from docx import Document

BASE_DIR = "generated"

def create_cover_letter_docx(filename, content):
    folder = os.path.join(BASE_DIR, "cover_letters")
    os.makedirs(folder, exist_ok=True)

    file_path = os.path.join(folder, filename)

    doc = Document()
    doc.add_heading("Cover Letter", level=1)
    doc.add_paragraph(content)   
    doc.save(file_path)

    return file_path

def create_resume_docx(filename, content):
    folder = os.path.join(BASE_DIR, "resumes")
    os.makedirs(folder, exist_ok=True)

    file_path = os.path.join(folder, filename)

    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph(content)
    doc.save(file_path)

    return file_path
